"""Exportação de gerações para .docx."""
import io
from datetime import datetime


def generation_to_docx(generation) -> bytes:
    """Converte uma Generation em bytes de arquivo .docx."""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Margens
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3)
        section.right_margin  = Cm(2.5)

    # Cabeçalho — título do documento
    title = doc.add_heading(generation.titulo or generation.tool_label, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.runs[0]
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    # Metadata
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.LEFT
    mr = meta.add_run(
        f"Gerado em {generation.created_at.strftime('%d/%m/%Y às %H:%M')} "
        f"· ContaIA — Ferramenta: {generation.tool_label}"
    )
    mr.font.size = Pt(9)
    mr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    mr.italic = True

    # Linha separadora (parágrafo vazio com borda inferior)
    sep = doc.add_paragraph()
    sep.paragraph_format.space_after = Pt(12)

    # Conteúdo gerado
    for linha in (generation.resultado or "").split("\n"):
        p = doc.add_paragraph()
        if linha.startswith("# "):
            run = p.add_run(linha[2:])
            run.bold = True
            run.font.size = Pt(13)
        elif linha.startswith("## "):
            run = p.add_run(linha[3:])
            run.bold = True
            run.font.size = Pt(11)
        elif linha.strip() == "":
            p.paragraph_format.space_after = Pt(4)
        else:
            run = p.add_run(linha)
            run.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(6)

    # Rodapé
    doc.add_paragraph()
    rodape = doc.add_paragraph()
    rodape.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = rodape.add_run("Documento gerado pelo ContaIA · Revise antes de enviar.")
    rr.font.size = Pt(8)
    rr.font.color.rgb = RGBColor(0xaa, 0xaa, 0xaa)
    rr.italic = True

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
